"""
Generate the final Internship Report in Word format.
Usage: python generate_report.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = r"f:\Users\30306\Desktop\副业\小雷\Y2742-白俄罗斯系统+实习报告（每周一次，总共四次）+论文-开发费：1300"

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_paragraph(doc, text):
    return doc.add_paragraph(text)

def add_bold_paragraph(doc, label, text):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    p.add_run(text)
    return p

def create_report(output_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_heading('Graduate Internship Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_sub.add_run("Belarus Tourism Review & Analysis System Development")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1A, 0x5E, 0x63)

    p_period = doc.add_paragraph()
    p_period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_period.add_run("Internship Period: June 1 – June 28, 2026")

    doc.add_paragraph()

    # 1. Introduction
    add_heading(doc, "1. Introduction", 1)
    add_paragraph(doc,
        "This report documents the development of the Belarus Tourism Review & Analysis System, "
        "a full-stack web application completed during a four-week graduate internship. The system "
        "was designed and built for Belarusian users to browse tourist attractions, write and read "
        "reviews, and receive AI-powered personalized recommendations. The entire project was "
        "implemented in English to meet the client's international requirements.")

    add_paragraph(doc,
        "The internship provided an opportunity to apply software engineering knowledge in a "
        "real-world project context, including full-stack development, AI integration, internationalization, "
        "and documentation practices.")

    # 2. Project Background
    add_heading(doc, "2. Project Background and Requirements", 1)
    add_paragraph(doc,
        "The client required a tourism platform tailored for Belarus with the following core features:")

    items = [
        "Attraction browsing and management (list, detail, search, filter)",
        "Travel review and diary publishing with star ratings",
        "AI-powered sentiment analysis on all reviews (positive / neutral / negative)",
        "Personalized attraction recommendations based on user rating history",
        "Rating prediction for unvisited attractions",
        "Full English user interface with internationalization support",
        "Admin panel for attraction management"
    ]
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    add_paragraph(doc,
        "The system was developed using Flask (Python) for the backend REST API and Vue 3 for the "
        "frontend, with SQLite as the database. AI features were powered by the DeepSeek API.")

    # 3. System Architecture
    add_heading(doc, "3. System Architecture", 1)
    add_paragraph(doc,
        "The application follows a client-server architecture with a clear separation between the "
        "frontend and backend layers.")

    add_heading(doc, "3.1 Backend (Flask)", 2)
    add_paragraph(doc,
        "The Flask backend provides a RESTful API with the following components:")

    backend_items = [
        "Flask-SQLAlchemy ORM with SQLite database",
        "Flask-JWT-Extended for JWT-based authentication",
        "Flask-CORS for cross-origin frontend access",
        "DeepSeek API integration for sentiment analysis and recommendations",
        "RESTful routes for auth, attractions, reviews, and recommendations"
    ]
    for item in backend_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    add_paragraph(doc,
        "Data models: User, Attraction, Review, Rating. The Review model includes sentiment "
        "analysis fields (sentiment_label, sentiment_score) computed via DeepSeek API.")

    add_heading(doc, "3.2 Frontend (Vue 3)", 2)
    add_paragraph(doc,
        "The Vue 3 frontend is a single-page application (SPA) with the following features:")

    frontend_items = [
        "Vue Router for client-side routing",
        "vue-i18n for internationalization (English only)",
        "Axios with JWT interceptors for API communication",
        "Responsive CSS design with teal/green theme",
        "Pages: Home, Attractions List, Attraction Detail, Login, Register, Recommendations, Profile"
    ]
    for item in frontend_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    add_heading(doc, "3.3 AI Integration", 2)
    add_paragraph(doc,
        "DeepSeek API (model: deepseek-chat) is used for two purposes:")
    ai_items = [
        "Sentiment Analysis: Each submitted review is analyzed to produce a label (positive/neutral/negative) "
        "and a score from -1.0 to 1.0. Results are stored in the database and displayed to users.",
        "Personalized Recommendations: Based on the user's rating history, DeepSeek generates "
        "a list of recommended attractions with reasons for each recommendation.",
        "Rating Prediction: DeepSeek estimates how a user would rate an unvisited attraction "
        "based on their past rating behavior."
    ]
    for item in ai_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    add_paragraph(doc,
        "A keyword-based fallback is implemented in case the DeepSeek API is unavailable.")

    # 4. Technical Implementation
    add_heading(doc, "4. Technical Implementation", 1)

    add_heading(doc, "4.1 Database Schema", 2)
    add_paragraph(doc,
        "The SQLite database contains four tables: users (username, email, password_hash, role), "
        "attractions (name, description, location, image_url, avg_rating, sentiment_score, review_count), "
        "reviews (user_id, attraction_id, content, rating, sentiment_label, sentiment_score), "
        "and ratings (user_id, attraction_id, score). Unique constraints prevent duplicate ratings.")

    add_heading(doc, "4.2 API Endpoints", 2)
    api_items = [
        "POST /api/auth/register — User registration",
        "POST /api/auth/login — User login, returns JWT token",
        "GET /api/attractions — List attractions (pagination, search, sort)",
        "GET/POST /api/attractions/:id — Attraction detail or creation",
        "GET/POST /api/reviews — List/create reviews",
        "GET /api/recommendations — Personalized recommendations (requires auth)",
        "POST /api/recommendations/predict-rating — Predict user rating for an attraction"
    ]
    for item in api_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    add_heading(doc, "4.3 Frontend Views", 2)
    view_items = [
        "Home: Hero banner, featured attractions, recent reviews",
        "Attractions List: Search, filter by location, sort by rating/name/date, pagination",
        "Attraction Detail: Full info, sentiment bar, review submission with star rating, predicted rating",
        "Login / Register: JWT-based authentication forms",
        "Recommendations: AI-generated personalized attraction suggestions",
        "Profile: User info and personal review history"
    ]
    for item in view_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    # 5. Weekly Summary
    add_heading(doc, "5. Weekly Development Summary", 1)

    weeks = [
        ("Week 1 (June 1–7)", "Set up project structure, designed database schema, implemented Flask models and auth API. Resolved environment configuration issues."),
        ("Week 2 (June 8–14)", "Built attractions and reviews APIs, integrated DeepSeek API for sentiment analysis, implemented rating prediction and recommendation endpoints. Started Vue 3 frontend scaffolding."),
        ("Week 3 (June 15–21)", "Completed all Vue 3 pages, integrated frontend with backend APIs, implemented sentiment visualization, applied responsive CSS styling, verified end-to-end functionality."),
        ("Week 4 (June 22–28)", "Compiled documentation: 4 Practice Diary reports and this Internship Report. Performed final system verification and prepared deployment package.")
    ]
    for week, desc in weeks:
        p = doc.add_paragraph()
        p.add_run(f"{week}: ").bold = True
        p.add_run(desc)

    # 6. Challenges
    add_heading(doc, "6. Challenges and Solutions", 1)
    challenges = [
        ("AI API Reliability:", "DeepSeek API occasional timeouts were handled by implementing a keyword-based fallback sentiment analyzer that activates when the API is unavailable."),
        ("Dependency Conflicts:", "Python environment had conflicting package versions. Resolved by installing packages in a clean environment and using specific version constraints in requirements.txt."),
        ("Internationalization:", "Ensuring all UI text was in English required reviewing every component and using vue-i18n consistently throughout the application."),
        ("Frontend-Backend Integration:", "Cross-origin requests were managed via Flask-CORS. JWT tokens are attached automatically via axios interceptors.")
    ]
    for label, text in challenges:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(text)

    # 7. Skills Developed
    add_heading(doc, "7. Skills and Knowledge Gained", 1)
    skills = [
        "Full-stack web development with Flask and Vue 3",
        "RESTful API design and implementation",
        "SQLAlchemy ORM and database schema design",
        "JWT-based authentication and security best practices",
        "Integration of LLMs (DeepSeek) for NLP tasks (sentiment analysis)",
        "AI prompt engineering for recommendation systems",
        "Internationalization (i18n) in Vue applications",
        "Responsive frontend design with modern CSS",
        "Software documentation and technical reporting"
    ]
    for skill in skills:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(skill)

    # 8. Conclusion
    add_heading(doc, "8. Conclusion", 1)
    add_paragraph(doc,
        "The Belarus Tourism Review & Analysis System was successfully developed within the four-week "
        "internship period. The system meets all client requirements: English interface, attraction "
        "management, review publishing with AI sentiment analysis, personalized AI recommendations, "
        "and rating prediction. The project demonstrates the practical application of full-stack "
        "development skills and AI integration in a real-world context.")

    add_paragraph(doc,
        "Future improvements could include adding user profile pictures, implementing a more advanced "
        "collaborative filtering recommendation engine, and deploying the application to a cloud "
        "platform for public access.")

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == '__main__':
    path = f"{BASE_DIR}\\InternshipReport.docx"
    create_report(path)
    print("Internship Report generated.")
