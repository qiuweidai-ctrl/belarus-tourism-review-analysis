"""
Generate Practice Diary (Week 1-4) in Word format.
Usage: python generate_diary.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE_DIR = r"f:\Users\30306\Desktop\副业\小雷\Y2742-白俄罗斯系统+实习报告（每周一次，总共四次）+论文-开发费：1300"

DIARIES = [
    {
        "week": 1,
        "date_range": "June 1 – June 7, 2026",
        "project": "Belarus Tourism Review & Analysis System",
        "content": {
            "project_overview": "Started working on the Belarus Tourism Review & Analysis System, a full-stack web application built with Flask (Python) for the backend and Vue 3 for the frontend. The system is designed for Belarusian users to browse attractions, write reviews, and receive AI-powered recommendations. The entire user interface is in English to meet the project requirements.",
            "work_done": [
                "Analyzed project requirements and created a detailed development plan",
                "Set up project directory structure with backend (Flask) and frontend (Vue 3) folders",
                "Designed database schema with SQLite: User, Attraction, Review, Rating tables",
                "Implemented Flask backend with SQLAlchemy ORM models",
                "Set up Flask-CORS and Flask-JWT-Extended for API authentication",
                "Created REST API endpoints for user registration and login",
                "Configured development environment and verified dependencies"
            ],
            "problems": [
                "Initial uncertainty about project scope and required features – resolved by communicating with the client and reviewing the requirements document carefully",
                "Dependency conflicts in the Python environment – resolved by creating a clean virtual environment"
            ],
            "next_week": "Complete the backend API for attractions and reviews, integrate DeepSeek API for sentiment analysis, and start the Vue 3 frontend development"
        }
    },
    {
        "week": 2,
        "date_range": "June 8 – June 14, 2026",
        "project": "Belarus Tourism Review & Analysis System",
        "content": {
            "project_overview": "Continued development of the Belarus Tourism Review & Analysis System. Focused on completing the backend API layer including attraction management, review submission, and implementing AI-powered sentiment analysis using the DeepSeek API. Also started frontend scaffolding with Vue 3.",
            "work_done": [
                "Implemented attraction management API (list, create, update, delete)",
                "Built review submission API with automatic sentiment analysis via DeepSeek API",
                "Developed rating prediction feature based on user history and attraction statistics",
                "Implemented personalized recommendation endpoint using DeepSeek AI",
                "Added pagination, search, and sorting to the attractions listing API",
                "Created Vue 3 project structure with Vue Router and vue-i18n for internationalization",
                "Designed frontend components and page layouts for all major views",
                "Initialized SQLite database with 6 sample Belarusian attractions"
            ],
            "problems": [
                "DeepSeek API occasional timeout issues – implemented fallback keyword-based sentiment analysis",
                "Watchdog import error with Flask debug mode – resolved by disabling debug hot-reload"
            ],
            "next_week": "Complete all Vue 3 frontend pages, connect frontend to backend APIs, perform end-to-end testing, and prepare the first Practice Diary report"
        }
    },
    {
        "week": 3,
        "date_range": "June 15 – June 21, 2026",
        "project": "Belarus Tourism Review & Analysis System",
        "content": {
            "project_overview": "This week focused on completing the Vue 3 frontend, implementing all UI pages, and conducting end-to-end integration testing. The frontend features English-only internationalized UI, responsive design, and full integration with the Flask backend APIs. The AI sentiment analysis and recommendation features were also tested.",
            "work_done": [
                "Completed all Vue 3 pages: Home, Attraction List, Attraction Detail, Login, Register, Recommendations, Profile",
                "Implemented star-rating input, sentiment badge display, and sentiment bar visualization",
                "Integrated axios interceptors for JWT token authentication",
                "Applied responsive CSS styling with a teal/green color theme",
                "Configured Vite to proxy API requests to Flask backend",
                "Built and deployed frontend assets to Flask static folder",
                "Verified all API endpoints with integration tests",
                "Populated database with sample data and default admin/demo accounts"
            ],
            "problems": [
                "Pinia useStore import error during build – removed unused import from App.vue",
                "Frontend build output directory outside project root warning – resolved with proper vite config"
            ],
            "next_week": "Write the 4 Practice Diary reports covering the entire internship period, and begin drafting the final Internship Report"
        }
    },
    {
        "week": 4,
        "date_range": "June 22 – June 28, 2026",
        "project": "Belarus Tourism Review & Analysis System",
        "content": {
            "project_overview": "The final week focused on documentation: writing all 4 Practice Diary reports and the final Internship Report. All system features were reviewed and verified. The system is fully functional with English interface, AI sentiment analysis, personalized recommendations, and rating prediction capabilities.",
            "work_done": [
                "Compiled Practice Diary reports for Weeks 1–4 documenting all development work",
                "Reviewed and verified all system functionality: user auth, attractions, reviews, AI features",
                "Documented the system architecture, API endpoints, and deployment instructions",
                "Created README.md with setup instructions and feature overview",
                "Verified DeepSeek API integration for both sentiment analysis and recommendation",
                "Ensured all UI text is in English with vue-i18n internationalization support",
                "Tested system end-to-end with demo user accounts (admin/demo)"
            ],
            "problems": [
                "Ensuring consistent English terminology throughout all documentation – resolved by reviewing each report carefully",
                "Balancing technical detail and readability in the internship report – resolved by following the provided example template"
            ],
            "next_week": "Submit all Practice Diary reports and the final Internship Report to the client"
        }
    }
]


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_paragraph(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    return p


def create_diary(week_num, data, output_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_heading('Graduate Practice Diary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Week info
    p_week = doc.add_paragraph()
    p_week.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_week.add_run(f"Week {week_num}  |  {data['date_range']}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x5E, 0x63)

    doc.add_paragraph()

    # Project
    add_heading(doc, "1. Project Overview", 1)
    add_paragraph(doc, data['content']['project_overview'])

    # Work done
    add_heading(doc, "2. Work Completed This Week", 1)
    for item in data['content']['work_done']:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

    # Problems
    add_heading(doc, "3. Problems Encountered and Solutions", 1)
    for i, (prob, sol) in enumerate(zip(data['content']['problems'], [''] * len(data['content']['problems'])), 1):
        p = doc.add_paragraph()
        p.add_run(f"Problem {i}: ").bold = True
        p.add_run(prob)
        p = doc.add_paragraph()
        p.add_run(f"Solution {i}: ").bold = True
        p.add_run(sol)

    # Next week
    add_heading(doc, "4. Next Week's Plan", 1)
    add_paragraph(doc, data['content']['next_week'])

    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == '__main__':
    for d in DIARIES:
        path = f"{BASE_DIR}\\PracticeDiary_Week{d['week']}.docx"
        create_diary(d['week'], d, path)
    print("All Practice Diary reports generated.")
